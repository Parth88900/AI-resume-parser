import spacy
import pdfplumber
import re
from docx import Document
import nltk
import os
import subprocess
import sys

# Install spaCy model if not available
try:
    import spacy
    spacy.load("en_core_web_sm")
except (ImportError, OSError):
    print("Downloading spaCy model...")
    subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])

# Download NLTK data with error handling
try:
    nltk.download('stopwords', quiet=True)
    nltk.download('punkt', quiet=True)
except Exception as e:
    print(f"NLTK download warning: {e}")

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

class ResumeParser:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # If model not found, download it
            os.system("python3 -m spacy download en_core_web_sm")
            self.nlp = spacy.load("en_core_web_sm")
        
        # Initialize stopwords with fallback
        try:
            self.stop_words = set(stopwords.words('english'))
        except LookupError:
            # If stopwords still not available, use a basic set
            self.stop_words = {'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 
                             "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 
                             'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 
                             'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 
                             'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 
                             'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 
                             'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 
                             'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 
                             'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 
                             'with', 'about', 'against', 'between', 'into', 'through', 'during', 
                             'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 
                             'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 
                             'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 
                             'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 
                             'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 
                             's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 
                             'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 
                             'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', 
                             "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 
                             'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', 
                             "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 
                             'won', "won't", 'wouldn', "wouldn't"}
        
        # Enhanced skills database
        self.skills_db = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'go', 'rust', 'swift', 'kotlin'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'django', 'flask', 'spring', 'node.js', 'express'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'cassandra'],
            'devops': ['docker', 'kubernetes', 'jenkins', 'aws', 'azure', 'gcp', 'ci/cd', 'terraform', 'ansible'],
            'data_science': ['machine learning', 'deep learning', 'nlp', 'computer vision', 'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn'],
            'tools': ['git', 'jira', 'confluence', 'linux', 'bash', 'powershell']
        }
        
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file with better error handling"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    try:
                        # Try multiple extraction methods
                        page_text = page.extract_text()
                        if not page_text:
                            # Try alternative extraction
                            page_text = page.extract_text(x_tolerance=1, y_tolerance=1)
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Error extracting text from page: {e}")
                        continue
                        
            # If still no text, try extracting tables
            if not text.strip():
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    text += ' '.join([str(cell) for cell in row if cell]) + "\n"
                                    
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
        
        return text
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file with better extraction"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
                            
            return text
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
    
    def extract_text(self, file_path):
        """Extract text based on file type"""
        if file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
    
    def preprocess_text(self, text):
        """Clean and preprocess the text"""
        if not text:
            return ""
            
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.@-]', '', text)
        return text.strip()
    
    def extract_contact_info(self, text):
        """Extract contact information"""
        contact_info = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['email'] = emails[0] if emails else "Not found"
        
        # Phone numbers
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        contact_info['phone'] = phones[0] if phones else "Not found"
        
        # LinkedIn
        linkedin_pattern = r'(https?://)?(www\.)?linkedin\.com/(in|company)/[a-zA-Z0-9-]+'
        linkedin_matches = re.findall(linkedin_pattern, text)
        if linkedin_matches:
            contact_info['linkedin'] = ''.join(linkedin_matches[0])
        else:
            contact_info['linkedin'] = "Not found"
        
        return contact_info
    
    def extract_name(self, text):
        """Extract candidate name using NER"""
        if not text:
            return "Not found"
            
        doc = self.nlp(text[:1000])  # Process first 1000 characters for efficiency
        names = []
        
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                # Filter out common false positives
                if len(ent.text.split()) >= 2 and len(ent.text) > 3:
                    names.append(ent.text)
        
        return names[0] if names else "Not found"
    
    def extract_skills(self, text):
        """Extract skills using comprehensive skill database"""
        if not text:
            return {}
            
        text_lower = text.lower()
        skills_found = {}
        
        for category, skills in self.skills_db.items():
            category_skills = []
            for skill in skills:
                # Check for exact matches and variations
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    category_skills.append(skill)
            if category_skills:
                skills_found[category] = category_skills
                
        return skills_found
    
    def extract_education(self, text):
        """Extract education information"""
        if not text:
            return []
            
        education_keywords = [
            'bachelor', 'master', 'phd', 'mba', 'bs', 'ms', 'b.tech', 'm.tech',
            'bsc', 'msc', 'ba', 'ma', 'university', 'college', 'institute', 
            'degree', 'graduated', 'diploma'
        ]
        
        sentences = text.split('.')
        education_info = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in education_keywords):
                # Clean up the sentence
                clean_sentence = re.sub(r'\s+', ' ', sentence).strip()
                if len(clean_sentence) > 10:  # Avoid very short sentences
                    education_info.append(clean_sentence)
                
        return education_info[:5]  # Return top 5 education entries
    
    def extract_experience(self, text):
        """Extract work experience information"""
        if not text:
            return {'years': "Not specified", 'companies': []}
            
        experience_patterns = [
            r'(\d+)\s*years?\s*of?\s*experience',
            r'experience.*(\d+)\s*years?',
            r'worked\s*for\s*(\d+)\s*years?',
            r'(\d+)\+?\s*years?\s*in',
            r'(\d+)\+?\s*years?\s*professional'
        ]
        
        experience = {'years': "Not specified", 'companies': []}
        
        # Find years of experience
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                experience['years'] = f"{matches[0]} years"
                break
        
        # Extract potential company names using NER
        doc = self.nlp(text)
        companies = []
        for ent in doc.ents:
            if ent.label_ == "ORG" and len(ent.text) > 3:
                companies.append(ent.text)
        
        experience['companies'] = list(set(companies))[:5]  # Remove duplicates and limit to 5
        
        return experience
    
    def parse_resume(self, file_path):
        """Main method to parse resume"""
        try:
            # Extract text based on file type
            text = self.extract_text(file_path)
            
            if not text or len(text.strip()) < 50:
                # Try to get more debug info
                print(f"Debug: Extracted text length: {len(text) if text else 0}")
                if text:
                    print(f"Debug: First 100 chars: {text[:100]}")
                return {"error": "The document appears to be empty or too short. Please ensure it's a text-based PDF/DOCX (not scanned).", "success": False}
            
            # Preprocess text
            cleaned_text = self.preprocess_text(text)
            
            # Extract information
            parsed_data = {
                'name': self.extract_name(cleaned_text),
                'contact_info': self.extract_contact_info(cleaned_text),
                'skills': self.extract_skills(cleaned_text),
                'education': self.extract_education(cleaned_text),
                'experience': self.extract_experience(cleaned_text),
                'text_length': len(cleaned_text),
                'raw_text': cleaned_text[:1000] + "..." if len(cleaned_text) > 1000 else cleaned_text,
                'success': True
            }
            
            return parsed_data
            
        except Exception as e:
            return {"error": f"Error parsing resume: {str(e)}", "success": False}
